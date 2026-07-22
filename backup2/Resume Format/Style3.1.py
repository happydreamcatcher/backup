import os
from datetime import datetime
import time
import pyperclip
import pyautogui
import keyboard
from weasyprint import HTML, CSS
from docx import Document

# ================= CONFIG =================
SAVE_FILE = "save.docx"
PDF_FILE = "MyCV2026.pdf"
HTML_FILE = "CV.html"
GPT_APP_TITLE = "ChatGPT"  # Adjust to your app window title
PROMPT_TEXT = "Please improve this text: "

# ==========================

def save_to_word(text):
    """Append text to a Word file"""
    today = datetime.now().strftime("%m-%d")
    filename = "Description\\" + today + ".docx"
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)
    print(f"Saved to {filename}")


def send_to_chatgpt():
    """F2: Copy selected text, send it to ChatGPT for generating resume"""
    # Copy selected text (Job Description)
    keyboard.press_and_release('ctrl+c')
    time.sleep(0.2)
    text = pyperclip.paste()
    if not text.strip():
        print("No text selected. Please select text first.")
        return

    save_to_word(text)

    pyautogui.keyDown('ctrl')
    pyautogui.press('tab')
    pyautogui.keyUp('ctrl')

    # Open new chat
    pyautogui.hotkey('ctrl', 'shift', 'o')
    time.sleep(0.5)

    # Paste prompt + text for ChatGPT
    pyperclip.copy(text + """

Above is job description and following is current resume and I am tailoring my resume according to job description.
And I also want to make resume.
1. Resume
I am gonnna make resume in order to bid on jobs.
Here is my resume skeleton and rule of how to structure and make resume.
------* Resume Skeleton --------
0. Head
Peter Nguyen
Senior Software Engineer
ng.peter.sky@outlook.com | Słupsk, Poland | www.linkedin.com/in/peter-nguyen-4680793b7/ | github.com/happydreamcatcher

1. Summary
2. Skills
3. Experience
- Accenture | Senior Software Engineer | Aug 2024 - Present
- Novo Nordisk	| Senior Software Engineer | Apr 2021 - Aug 2024		              
- Deloitte | Application Engineer | Dec 2018 - Mar 2021 	              
- Objectivity Inc | Web Developer | Jul 2017 - Nov 2018                      
4. Education
  Hanoi University of Science and Technology | Bachelor's Degree in Computer Programming | Hanoi, Vietnam | 2013 - 2017 		  

------* Resume make rules --------
1. Summary Section
Write a strong professional summary(8 years).
In 4 concise sentences, summarize my experience, core strengths, and what makes me a good fit for the role. 
Showcase to my extensive experience including the project type and project names briefly that I have been participated in last three companies(Accenture, Novo Nordisk, Deloitte, Objectivity.Inc: Never mention company names).
When next sentence starts in same line with previous sentence, leave 4 letters gap between sentences in summary.
 
2. Core Skills Section
 - At least, skill section must include almost tech stack or skills of JD.
   Excepted for JD mentioned Skills, also, plz  understand what related skills are essential or necessary and, mention them even though they are not stated in JD.
 - (skill section should include 25 - 35 skills)
 - under category, give me following visual way:
   Category1 Name(bold font): Skill1, Skill2(all in one line)
   Category1 Name(bold font): Skill1, Skill2((all in one line))
   Use bullet point for every category.
   don't hightlight keywords at all.
 - The maximum category numbers are 8. Never make more than 8 categories.
3. Experience Section
  1)Project overview & Scope in every company 
   - Accenture: Senior Role
	 * Accenture is software dev company. so I can work various type of projects. After reading JD, make very specific project name(less popular and small project) in a related industry and definitely different project with very reasonable sceario.
     Here, most important thing is that project name must not sound like generic. It must sound real world project name and scenario too. I hate ("IAM system for Sass platform") like generic project name. When hiring manager sees the project name, he must understand oh what this project is for.
     Also, project name must not include platform, it sounds awkward, make sense?
   - Novo Nordisk: Senior Role
   * Novo Nordisk is a global healthcare company headquartered in Denmark, focused on treatments and care solutions for diabetes, obesity, and other serious chronic diseases. In addition to medicines, the company invests in digital health, connected care, and patient support solutions to improve disease management and real-world outcomes.
   * Work Scope:  
     •	Digital Patient Support & Treatment Engagement Platform: Worked as part of the development team on patient-facing digital health initiatives designed to support people living with chronic diseases such as diabetes and obesity. Contributed to web experiences focused on treatment education, onboarding, ongoing support, and engagement features that helped patients better manage therapy and stay on prescribed treatment plans. This aligns with Novo Nordisk’s stated focus on digital therapeutics, patient support solutions, treatment initiation, adherence, and chronic disease management.
     •	Connected Care & HCP-Facing Disease Management Portal: Contributed to digital platform work supporting connected care and healthcare-professional workflows, including web interfaces for chronic disease education, therapy-related guidance, and integration-oriented experiences tied to Novo Nordisk’s broader digital health ecosystem. The project scope fits Novo Nordisk’s work around digital partnerships, developer tools, connected solutions, and HCP-focused disease education, as well as smart-device/app ecosystems in diabetes care.
  - Deloitte: Mid role
   * Deloitte is a global professional services and consulting firm that delivers strategy, technology, engineering, cloud, data, cybersecurity, and industry-specific transformation solutions for enterprise and public sector clients. Its engineering, AI, and data practice focuses on modernizing mission-critical platforms and digital operations, while its consulting and risk capabilities also include solutions such as digital identity and transformation delivery platforms.
   * Work Scope:  
     •	Enterprise Identity & Access Governance Portal: Worked as part of the delivery team for a large enterprise client on a digital identity and access management initiative, contributing to web-based administrative and user-facing workflows for role-based access requests, approval routing, password/self-service flows, and identity lifecycle operations. This fits Deloitte’s stated digital identity offering around integrating and operating the identity lifecycle under one managed solution.
     •	Digital Transformation Program Delivery Dashboard: Contributed to a transformation-management platform for an enterprise client, building web features that helped teams track delivery milestones, dependencies, risks, and execution progress across multiple workstreams. The scope aligns with Deloitte’s Intelligent Delivery Platform, which is designed to manage and connect digital transformation initiatives while streamlining program delivery.
	 * Never say : played the key role, or lead ... or mentor other guys. I started as software intern then promoted to software developer in Miquido.
   - Objectivity.Inc Early starting career
   * Objectivity was a digital engineering and software consulting firm focused on cloud-native platforms, product engineering, and data modernization for clients across financial services, healthcare, retail, manufacturing, and smart office domains.
   * Work Scope:  
     •	Financial Services Workflow Modernization Platform: Contributed to the development of a cloud-native web platform for a financial services client, supporting case-based operations, internal workflow modernization, and improved access to operational data across legacy-connected systems.
     •	Connected Workplace Experience Portal: Worked on web-based product features for a smart office client, helping deliver connected workplace workflows and platform functionality that improved employee interaction with office services and digital operations.
  2)Experience making rule
 - I gave you project name or scope or scenario for every company above.
  In Next step, while relying on the project scope and company overview which I gave you above , also you must try to  tailore to JD so that I could be best candidate for this position(passing over 95% ATS).  
   plz read JD carefully and understand what responsibility(and ability)  this role require me to have. then try to make resume so that I have perfectly suitable ability for this role. 
   I mean plz try to combine my project scope or scenario with JD's role sealessly very very naturally.
   (JD's requirement <- Balance -> Project Scope)
 - If role is backend, or frontend or fullstack or data engineering, make relevant experience solely for corresponding field
   (I mean JD's role is backend, then must not mention frontend, also if JD mention data engineernig, don't mention others, but only mention data engineering
    Please comply with this rule for all companies) 
 - Sometimes, job description does not include enough skills and experience and they mention only a few core  skills or experience rather than listing many.
   So plz think what additional skills and experience would be good to be added and plz add them in experience section. 
   (ex: React is mentioned in JD, but Next.js is not mentioned. in this case You should mention Next.js in skill section too. This is vivid example of mentioning related skills.)
 - make 7 for Accenture, 7 bullet work bio for Novo Nordisk, 6 for Deloitte, 6 for Objectivity.Inc.
 - as long as possible technically, try to mention many core and important skills and keywords mentioned in JD when making experience section bullet points.
   and also make sentences long professionally.
 - Never make technically fake sentences.
 - IF JD does not mention much enough skills/tools/experience, then plz imagine what are necessary then mention them. 
 - 70% focus on JD, 30% focus on supporting with additional skills/tools/experience.
 - hightlight(bold font) as many skills as possible in work experience to showcase my experience of I have already used them.
 - end each bullet with fullstop.
 - Never use "Led ..." if you want to emphasize that kind of leadership skill, use "Played the key role in ... "
   Never use these expressions: "scalable", "high-performance".                
 3)Quantifiable Impact - Non-Negotiable:
  * MUST add quantifiable results in 2–3 technical bullet points per company, using strict numeric expressions (e.g., “around X records,” “over X users,” “approximately X requests per day”).
    * **NEVER show quantifiable impacts in every bullet point**. Only 1 -2 bullets should feature quantifiable metrics that reflect the most significant impact of your work. The other bullets should focus on describing actions, responsibilities, or broader context without quantification.
  * Quantitative values should be determined based on the company stage and available data. For early-stage or smaller companies, provide realistic approximations of scale. For more established companies, use exact numbers or well-supported approximations based on available metrics.
  * In most cases, if precise figures are unavailable, MUST use approximations like “around,” “over,” or “approximately” to convey a sense of scale or impact.
    * Do not use vague, non-numeric descriptors or qualitative terms.
  * Quantification may include:
    * Exact numeric volumes (e.g., records, requests)
    * Numeric user counts (e.g., user base, active users)
    * Numeric timeframes (e.g., release cycles, project duration)
    * Team or system scales (e.g., number of services, team size)
    * Workload or throughput values
    * Deployment frequencies (e.g., updates, releases per time period)
  * MUST NOT use:
    * Textual numeric expressions (e.g., *thousands, hundreds, tens of thousands*)
    * Plus symbols (e.g., *X+ users*)
    * Percentages in any form
    * Vague, non-numeric descriptors (e.g., *a lot of*, *many*, *significant*).
  **Important (Required for All Bullet Points)**: NEVER use percentage expressions like "x%", MUST remove or replace it with numeric expressions.

Please make html in new code panel with following styles:
I give you sample html code for one resume below:
                   
<!DOCTYPE html>

<html lang="en">
<head>

  <meta charset="UTF-8">

  <meta name="viewport" content="width=device-width, initial-scale=1.0">

  <style>

    body {

      font-family: Georgia;

    }
    
    p {
      padding: 0;
      margin: 0;        
      line-height: 1.4;
    }
     

    .title {

      font-size: 16px;

      color: black;

    }



    .role {

      font-size: 14px;

    }



    .section-title {

      font-size: 14px;

      font-weight: bold;

      color: #6b6866;

      display: inline-block;

      padding-bottom: 5px;

    }



    .underline {

      width: 100%;

      height: 2px;

      background-color: black;

    }



    .content {

      font-size: 12px;

    }



    .bold {

      font-weight: bold;
      margin: 5px 0 5px 0;

    }



    ul {

      list-style-type: disc;

      padding-left: 20px;
      margin: 0;

    }
                   
    ul li {
      margin-bottom: 0px;     
      line-height: 1.4;      
    }                   

    .education-content {

      font-weight: bold;
      display: block;
      margin: 5px 0 5px 0;

    }

  </style>

</head>



<body>



  <!-- Header Section -->

  <div class="title">Michal Truong</div>

  <div class="role">Senior Software Engineer</div>

  <ul class="content" style="margin-top: 4px;"><li>${email}</li><li>${location}</li><li>${linkedin}</li><li>${github}</li></ul>



  <!-- Summary Section -->

  <div class="section-title" style="margin-top: 5px;">Summary</div>

  <div class="underline"></div>

  <div class="content">

    <p>Experienced JD's job title ....</p>

  </div>



  <!-- Skills Section -->

  <div class="section-title">Skills</div>

  <div class="underline"></div>

  <div class="content">

    <ul>

      <li><strong>[Category1]:</strong>[skill1], [skill2], [skill3], ...</li>

      <li><strong>[Category2]:</strong>[skill1], [skill2], [skill3], ...</li>
      (....)
    </ul>

  </div>



  <!-- Work Experience Section -->

  <div class="section-title">Work Experience</div>

  <div class="underline"></div>

  <div class="content">
    <!-- for all companies -->
    <div class="bold">${company name} • ${position title} • ${Period}</div>

    <ul>
      <li>[experience1.... <strong>${core skill},.... ${core skill}...]</strong></li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience3.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...
     </ul>

  </div>

  <!-- Education Section -->

  <div class="section-title">Education</div>

  <div class="underline"></div>

  <div class="content">

    <ul>

      <li><span class="education-content">Ho Chi Minh City University of Technology • Bachelor's Degree • Computational and Applied Mathematics • Ho Chi Minh City, Vietnam • 2012 - 2016</span></li>

    </ul>

  </div>



</body>



</html>
(Don't add another margin or padding in your decision)   

All in all, you must give me in code panle for resume so that I can copy.         
But before give me resume, gpt, plz judge whether job description say whether this job require me to work hybrid or on site.
If this is not fully remote, stop making resume and let me know the reason.(I am based in Poland).
If full remote, then go to next step.

Next step: As I emphasize much about Accenture, you must make very very specific experience section not generic sound one. but you always forget this critical rule cuz you are foolish.
So before making html resume, let me know your plan of project name & project scenario for  briefly(Think you are presenting your project idea in very detail with real word usage example). 
Here, important thing is that one engineer can not do everything for one project development, so never say that I participated in so many works. all bullets of experience must be specific while related to project name so you must make  hiring manger 
think, oh it is right and reasonable that this guy has worked on this project, not fake engineer. make sense? make sense? 

After briefly showing your plan for Accenture, also briefly mention what languages & frameworks you are gonna use for another companies briefly, Avoid to mention skills that are not explicitly mentioned in JD.
Also, I said you must use strong tag in html to make bold text not --, so after you add plan, plz confirm you will only use strong tag for bold text.
If I agree with you, 
then, start to make resume.                   
""")
    pyautogui.hotkey('ctrl', 'v')
    pyautogui.press('enter')
    print("Text sent to ChatGPT. Save DOC output manually as 'latest.docx', then press F3.")


def copy_selected_html_to_resume():
    """F3: Copy selected HTML code (resume), create resume.html, generate resume.pdf"""
    # Copy selected HTML from clipboard
    keyboard.press_and_release('ctrl+c')
    time.sleep(0.2)
    html_text = pyperclip.paste()
    if not html_text.strip():
        print("No HTML selected. Please select your HTML code first.")
        return
    
    # Append to save.docx
    save_to_word(html_text)
    print(f"Appended selected HTML to {SAVE_FILE}")
    
    # Create new resume.html (overwrite old if exists)
    if os.path.exists(HTML_FILE):
        os.remove(HTML_FILE)
    with open(HTML_FILE, "w", encoding="utf-8") as f:
        f.write(html_text)
    print(f"Created new {HTML_FILE}")

    # Generate PDF from resume.html
    if os.path.exists(PDF_FILE):
        os.remove(PDF_FILE)
	


    HTML(HTML_FILE, encoding="utf-8").write_pdf(
        PDF_FILE,
        stylesheets=[
            CSS(string="""
                @page {
                    size: A4;
                    margin-top: 10mm;
                    margin-bottom: 7mm;
                    margin-left: 5mm;
                    margin-right: 5mm;
                }
                body {
                    font-family: "Bell MT", sans-serif;
                    letter-spacing: 0.02em;
                }
            """)
        ]
    )

    os.startfile(PDF_FILE)
    print("Generated resume PDF!")

# ================= HOTKEYS =================
keyboard.add_hotkey('[', send_to_chatgpt)  # Press `[` to send Job Description to ChatGPT
keyboard.add_hotkey(']', copy_selected_html_to_resume)  # Press `]` to generate Resume PDF

print("[: Send selected job description to ChatGPT")
print("]: Generate Resume PDF from HTML")
keyboard.wait()
