import os
from datetime import datetime
import time
import pyperclip
import pyautogui
import keyboard
from weasyprint import HTML, CSS
from docx import Document

# ================= CONFIG =================
SAVE_FILE = "save.docx"
PDF_FILE = "DanielPhanCV.pdf"
HTML_FILE = "resume.html"
GPT_APP_TITLE = "ChatGPT"  # Adjust to your app window title
PROMPT_TEXT = "Please improve this text: "

# ==========================

def save_to_word(text):
    """Append text to a Word file"""
    today = datetime.now().strftime("%m-%d")
    filename = "Description\\" + today + ".docx"
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)
    print(f"Saved to {filename}")


def send_to_chatgpt():
    """F2: Copy selected text, send it to ChatGPT for generating resume"""
    # Copy selected text (Job Description)
    keyboard.press_and_release('ctrl+c')
    time.sleep(0.2)
    text = pyperclip.paste()
    if not text.strip():
        print("No text selected. Please select text first.")
        return

    save_to_word(text)

    pyautogui.keyDown('ctrl')
    pyautogui.press('tab')
    pyautogui.keyUp('ctrl')

    # Open new chat
    pyautogui.hotkey('ctrl', 'shift', 'o')
    time.sleep(0.5)

    # Paste prompt + text for ChatGPT
    pyperclip.copy(text + """

Above is job description and following is current resume and I am tailoring my resume according to job description.
And I also want to make resume.
1. Resume
I am gonnna make resume in order to bid on jobs.
Here is my resume skeleton and rule of how to structure and make resume.
------* Resume Skeleton --------
0. Head
Daniel Phan
Senior Software Engineer Engineer
 | Brasov, Romania | www.linkedin.com/in/daniel-phan-b6aa25401/

1. Summary
2. Skills
3. Experience
- Thoughtworks | Senior Software Engineer | 05/2025 - Present
- Deloitte | Senior Software Consultant | 08/2022 - 04/2025		              
- Zendesk | Application Engineer | 10/2020 - 07/2022              
- Bluebik Global | Software Developer | 01/2018 - 09/2020
                                         
4. Education
Danang University | Bachelor's Degree in Computational and Applied Mathematics | Vietnam | 08/2013 - 05/2017 		  

------* Resume make rules --------
1. Summary Section
Write a strong professional summary(8 years).
In 4 concise sentences, summarize my experience, core strengths, and what makes me a good fit for the role. 
Showcase to my extensive experience including the project type and project names briefly that I have been participated in companies: Never mention company names).
Don't highlight keywords in summary section. 
2. Core Skills Section
 - At least, skill section must include almost tech stack or skills of JD.
   Excepted for JD mentioned Skills, also, plz  understand what related skills are essential or necessary and, mention them even though they are not stated in JD.
   example: as a React frontend job, you can mention following skills(even though some of belows skills are not mentioned in JD: just vivid example):
   Styling & UI Technologies: SASS, Tailwind CSS, Styled Components, Material UI, jQuery, AntD, ShadCN
   UI, Mantine, Chakra UI, Bootstrap
   This is only one example case. The purpose of this is to show hiring managers that I have extensive skills which is either mentioned in JD or not mentioned in JD and also in order to not show I directly tailor my resume to JD. Make sense?
 - (skill section should include 25 - 32 skills)
 - under category, give me following visual way:
   Category1 Name(bold font): Skill1, Skill2(all in one line)
   Category1 Name(bold font): Skill1, Skill2((all in one line))
   Use bullet point for every category.
   don't hightlight keywords at all.
 - The number of skills shouldn't be more than 35.
3. Experience Section
  1)Project overview & Scope in every company
  - Thoughtworks: Senior Role
	 * Thoughtworks is software dev company. so I can work various type of projects. After reading JD, do web search and find very similar project (less popular and small project) which was developed by Capgemini recently.(Never say generic project name like Saa Product.)
     Then based on that project scenario, say I have worked on it. Don't say I did everything for that project. 
   - Deloitte: Senior Role
     Deloitte is large scale global consulting company including software field. so I can work various type of projects. 
    After reading JD, do web search and find very similar project (less popular and small project) which was developed by Thoughtworks recently.(Never say generic project name like Saa Product.)
     Then based on that project scenario, say I have worked on it. Don't say I did everything for that project.
     Be cautious: Should be different from Thoughtworks's scenario. But qualified for JD's role.
  - Zendesk: Mid role
   * Zendesk Overview:
     Zendesk provides a cloud-based (SaaS) customer service platform that helps businesses manage customer support across multiple channels like email, chat, phone, and social media.
     The company is currently focusing heavily on AI-powered tools to boost efficiency. Key features include:
     Omnichannel Support: A unified workspace for agents to handle requests from any channel.
     AI & Automation: Built-in AI for triage, routing, and proactive suggestions (Copilot).
     Self-Service Tools: Knowledge base and help center functionalities for customers.
     Extensive Integrations: Over 1,200+ apps and integrations in the Zendesk Marketplace, including deep ties with Salesforce and Twitter
   * Work Scope:  
     Based on company overview above, make experience section scenario.(I mentioned all work field done by Zendesk. But for resume, use only one field)
	 * Never say : played the key role, or lead ... or mentor other guys. I started as software intern then promoted to software developer in Miquido.
  - Bluebik Global: Junior  level role
   * Bluebik Global is a Thailand-based management consulting and digital transformation firm delivering strategy, data/AI, and software implementation for enterprise clients.
     It has local branch in Vietnam.
   * Work Scope:  
     •	On “Bank Customer 360 & Consent Platform,” I built backend services for customer profiles, consent/audit trails, and integrations with core banking/CRM systems.
     •  Supported the data side by helping model key datasets and delivering repeatable extracts for analytics teams to track campaign performance and customer segmentation.               
  2)Experience making rule
 - I gave you project name or scope or scenario for every company above.
  In Next step, while relying on the project scope and company overview which I gave you above , also you must try to  tailore to JD so that I could be best candidate for this position(passing over 95% ATS).  
   plz read JD carefully and understand what responsibility(and ability)  this role require me to have. then try to make resume so that I have perfectly suitable ability for this role. 
 - Sometimes, job description does not include enough skills and experience and they mention only a few core  skills or experience rather than listing many.
   So plz think what additional skills and experience would be good to be added and plz add them in experience section. 
   (ex: React is mentioned in JD, but Next.js is not mentioned. in this case You should mention Next.js in skill section too. This is vivid example of mentioning related skills.)
 - make 8 bullets for Thoughtworks, 7 for Deloitte, 6 for Zendesk or 6 for Bluebik Global
 - as long as possible technically, try to mention many core and important skills and keywords mentioned in JD when making experience section bullet points.
    and make experience section sentence long professionally.
 - Never make technically fake sentences.
 - When describing the experience, must focus about 70% on the primary programming languages, technology stacks, libraries, frameworks, and tools listed in the job description, and use the remaining 30% to highlight other related relevant technologies or skills that strengthen my overall profile. 
   You should mention core language or framework related another libraries or framework too. 
 - hightlight(bold font) as many skills as possible in work experience to showcase my experience of I have already used them.
 - Never use these expressions: "scalable", "high-performance".    
 - Never wrap project name with "".            
 3)Quantifiable Impact - Non-Negotiable:
  * MUST add quantifiable results in 2–3 technical bullet points per company, using strict numeric expressions (e.g., “around X records,” “over X users,” “approximately X requests per day”).
    * **NEVER show quantifiable impacts in every bullet point**. Only 1 -2 bullets should feature quantifiable metrics that reflect the most significant impact of your work. The other bullets should focus on describing actions, responsibilities, or broader context without quantification.
  * Quantitative values should be determined based on the company stage and available data. For early-stage or smaller companies, provide realistic approximations of scale. For more established companies, use exact numbers or well-supported approximations based on available metrics.
  * In most cases, if precise figures are unavailable, MUST use approximations like “around,” “over,” or “approximately” to convey a sense of scale or impact.
    * Do not use vague, non-numeric descriptors or qualitative terms.
  * Quantification may include:
    * Exact numeric volumes (e.g., records, requests)
    * Numeric user counts (e.g., user base, active users)
    * Numeric timeframes (e.g., release cycles, project duration)
    * Team or system scales (e.g., number of services, team size)
    * Workload or throughput values
    * Deployment frequencies (e.g., updates, releases per time period)
  * MUST NOT use:
    * Textual numeric expressions (e.g., *thousands, hundreds, tens of thousands*)
    * Plus symbols (e.g., *X+ users*)
    * Percentages in any form
    * Vague, non-numeric descriptors (e.g., *a lot of*, *many*, *significant*).
  **Important (Required for All Bullet Points)**: NEVER use percentage expressions like "x%", MUST remove or replace it with numeric expressions.
  ** Never make certification section.
Please make html in new code panel with following styles:
I give you sample html code for one resume below:
                   
<!DOCTYPE html>

<html lang="en">



<head>

  <meta charset="UTF-8">

  <meta name="viewport" content="width=device-width, initial-scale=1.0">

  <style>

    body {

      font-family: Bell MT, Helvetica, sans-serif;

    }



    .title {

      font-size: 16px;

      color: black;

    }



    .role {

      font-size: 14px;

    }



    .section-title {

      font-size: 14px;

      font-weight: bold;

      color: #3464f3;

      display: flex;
                   
      justify-content: center;

      padding-bottom: 5px;

    }



    .underline {

      width: 100%;

      height: 1px;

      background-color: black;

      # margin-bottom: 20px;

    }



    .content {

      font-size: 12px;

    }



    .bold {

      font-weight: bold;

    }



    ul {

      list-style-type: disc;
      line-height: 17px;
      margin-top: 0;
      padding-left: 20px;

    }



    li {

  

    }



    .education-content {

      font-weight: bold;

    }

  </style>

</head>



<body>



  <!-- Header Section -->

  <div class="title">Mateusz Wongchai</div>

  <div class="role">Senior Data Engineer</div>

  <div class="content" style="margin-top: 4px;">${email } • ${localtion } • ${linkedin link}</div>



  <!-- Summary Section -->

  <div class="section-title" style="margin-top: 5px;">SUMMARY</div>

  <div class="underline"></div>

  <div class="content">

    <p style="line-height: 17px;">Experienced JD's job title ....</p>

  </div>



  <!-- Skills Section -->

  <div class="section-title">SKILLS</div>

  <div class="underline"></div>

  <div class="content" style="margin-top: 5px;">

    <ul>

      <li><strong>[Category1]:</strong>[skill1], [skill2], [skill3], ...</li>

      <li><strong>[Category2]:</strong>[skill1], [skill2], [skill3], ...</li>
      (....)
    </ul>

  </div>



  <!-- Work Experience Section -->

  <div class="section-title">PROFESSIONAL EXPERIENCE</div>

  <div class="underline"></div>

  <div class="content">

    <!-- Across all companies -->

    <div class="bold" style="margin-top:10px;">${Company Name} | ${Position Title} | ${Period}</div>

    <ul>
      <li>[experience1.... <strong>${core skill},.... ${core skill}...]</strong></li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience3.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...
    </ul>
                   
  </div>



  <!-- Education Section -->

  <div class="section-title">EDUCATION</div>

  <div class="underline"></div>

  <div class="content" style="margin-top:10px;">

    <ul>

      <li><span class="education-content">Danang Univerisity of Science and Technology | Bachelor of Science, Computer Science | Vietnam | 08/2013 - 05/2017</span></li>

    </ul>

  </div>



</body>



</html>

All in all, you must give me in code panle for resume so that I can copy.         
But before give me resume, gpt, plz judge whether job description say whether this job require me to work hybrid or on site.
If this is not fully remote, stop making resume and let me know the reason.(I am based in Poland).
If full remote, then go to next step.

Next step: As I emphasize much about Thoughtworks, you must make very very specific experience section not generic sound one. but you always forget this critical rule cuz you are foolish.
So before making html resume, briefly mention what languages & frameworks & tools you are gonna mention across all companies companies.
If I agree with you plan, move to next step.
To avoid making awkward resume, you should mention necessary many standard data engineering tools/skills/exp even though these are not explicitly mentioned in JD.

Also, I said you must use strong tag in html to make bold text not --, so after you add plan, plz confirm you will only use strong tag for bold text.
If I agree with you, 
then, start to make resume.    
Caution: You sometimes use ** ** to make text bold(I mean highligh keywords...). This is banned. You must necessarily use only strong tag <strong> like above HTML file to make all bold text, especially in experience section.


                   
""")
    pyautogui.hotkey('ctrl', 'v')
    pyautogui.press('enter')
    print("Text sent to ChatGPT. Save DOC output manually as 'latest.docx', then press F3.")


def copy_selected_html_to_resume():
    """F3: Copy selected HTML code (resume), create resume.html, generate resume.pdf"""
    # Copy selected HTML from clipboard
    keyboard.press_and_release('ctrl+c')
    time.sleep(0.2)
    html_text = pyperclip.paste()
    if not html_text.strip():
        print("No HTML selected. Please select your HTML code first.")
        return
    
    # Append to save.docx
    save_to_word(html_text)
    print(f"Appended selected HTML to {SAVE_FILE}")
    
    # Create new resume.html (overwrite old if exists)
    if os.path.exists(HTML_FILE):
        os.remove(HTML_FILE)
    with open(HTML_FILE, "w", encoding="utf-8") as f:
        f.write(html_text)
    print(f"Created new {HTML_FILE}")

    # Generate PDF from resume.html
    if os.path.exists(PDF_FILE):
        os.remove(PDF_FILE)
	


    HTML(HTML_FILE, encoding="utf-8").write_pdf(
        PDF_FILE,
        stylesheets=[
            CSS(string="""
                @page {
                    size: A4;
                    margin-top: 10mm;
                    margin-bottom: 7mm;
                    margin-left: 5mm;
                    margin-right: 5mm;
                }
                body {
                    font-family: "Bell MT", sans-serif;
                    letter-spacing: 0.02em;
                }
            """)
        ]
    )

    os.startfile(PDF_FILE)
    print("Generated resume PDF!")

# ================= HOTKEYS =================
keyboard.add_hotkey('[', send_to_chatgpt)  # Press `[` to send Job Description to ChatGPT
keyboard.add_hotkey(']', copy_selected_html_to_resume)  # Press `]` to generate Resume PDF

print("[: Send selected job description to ChatGPT")
print("]: Generate Resume PDF from HTML")
keyboard.wait()
