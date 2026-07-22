import os
from datetime import datetime
import time
import pyperclip
import pyautogui
import keyboard
from weasyprint import HTML, CSS
from docx import Document

# ================= CONFIG =================
SAVE_FILE = "save.docx"
PDF_FILE = "CristianCV.pdf"
HTML_FILE = "resume.html"
GPT_APP_TITLE = "ChatGPT"  # Adjust to your app window title
PROMPT_TEXT = "Please improve this text: "

# ==========================

def save_to_word(text):
    """Append text to a Word file"""
    today = datetime.now().strftime("%m-%d")
    filename = "Description\\" + today + ".docx"
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)
    print(f"Saved to {filename}")


def send_to_chatgpt():
    """F2: Copy selected text, send it to ChatGPT for generating resume"""
    # Copy selected text (Job Description)
    keyboard.press_and_release('ctrl+c')
    time.sleep(0.2)
    text = pyperclip.paste()
    if not text.strip():
        print("No text selected. Please select text first.")
        return

    save_to_word(text)

    pyautogui.keyDown('ctrl')
    pyautogui.press('tab')
    pyautogui.keyUp('ctrl')

    # Open new chat
    pyautogui.hotkey('ctrl', 'shift', 'o')
    time.sleep(0.5)

    # Paste prompt + text for ChatGPT
    pyperclip.copy(text + """

Above is job description and following is current resume and I am tailoring my resume according to job description.
And I also want to make resume.
1. Resume
I am gonnna make resume in order to bid on jobs.
Here is my resume skeleton and rule of how to structure and make resume.
------* Resume Skeleton --------
0. Head
Cristian Chantarangsu
Senior Data Engineer
Bucharest, Romania • cristian.c.0808@hotmail.com • linkedin.com/in/cristian-chantarangsu-b5b931393/

1. Summary
2. Skills
3. Experience
- Fresenius Medical Care • Senior Data Engineer • 06/2023 - Present
- Capgemini	• Data Engineer • 09/2020 - 05/2023		              
- Novartis • Junior Software Developer • Poland • 05/2018 - 08/2020                      
4. Education
  Babes-Bolyai University • Bachelor's Degree in COMPUTER AND INFORMATION SCIENCES AND SUPPORT SERVICES • Romania • 2013 - 2017 		  

------* Resume make rules --------
1. Summary Section
Write a strong professional summary(8 years).
Start with "Experienced Senior Data Engineer..."
In 4 concise sentences, summarize my experience, core strengths, and what makes me a good fit for the role. 
Showcase to my extensive experience including the project type and project names briefly that I have been participated in last three companies(Fresenius Medical Care, Capgemini, Novartis: Never mention company names).
When next sentence starts in same line with previous sentence, leave 4 letters gap between sentences in summary.
 
2. Core Skills Section
 - At least, skill section must include almost tech stack or skills of JD.
   Excepted for JD mentioned Skills, also, plz  understand what related skills are essential or necessary and, mention them even though they are not stated in JD.
 - (skill section should include 25 - 35 skills)
 - under category, give me following visual way:
   Category1 Name(bold font): Skill1, Skill2(all in one line)
   Category1 Name(bold font): Skill1, Skill2((all in one line))
   Use bullet point for every category.
   don't hightlight keywords at all.
 - The maximum category numbers are 8. Never make more than 8 categories.
3. Experience Section
  1)Project overview & Scope in every company
   - Fresenius Medical Care:: Senior role
   * Company Explanation: Fresenius Medical Care is the world’s leading provider of products and services for people with renal diseases, including dialysis machines, consumables, and kidney care therapies through its global network of clinics and technology solutions.                  
   * Work Scope: Plz select one project below which is relevant to JD: 
    •	Apollo Global Clinical Data Platform: A harmonized cloud‑based database aggregating anonymized clinical data from hundreds of thousands of dialysis patients and millions of treatment records across dozens of countries to enable global analytics, quality monitoring, and clinical insight generation.
    •	Virtual & Augmented Reality Training Tools: Immersive AR and VR‑based training programs designed to educate healthcare professionals and patients on complex dialysis equipment and procedures, improving skill retention and treatment confidence.
    •	Digital Patient Engagement & Information Initiatives: Digital health efforts focused on connecting patients with their health information, enhancing patient portal capabilities, and exploring consumer‑directed tools for access, management, and exchange of healthcare data.
   - Capgemini: Mid-Senior level role
   * Capgemini is software dev company. so I can work various type of projects. After reading JD, make very specific project name(less popular and small project) in a related industry and definitely different project with very reasonable sceario.
     Here, most important thing is that project name must not sound like generic. It must sound real world project name and scenario too. I hate ("IAM system for Sass platform") like generic project name. When hiring manager sees the project name, he must understand oh what this project is for.
     Also, project name must not include platform, it sounds awkward, make sense?
   - Novartis: junior developer role  
   * Company Explanation: Novartis is a global innovative medicines company that researches, develops, manufactures, and markets healthcare products and therapies to improve and extend people’s lives around the world.                
   * Work Scope: Plz select one project below which is relevant to JD: 
    •	AI Innovation Lab & Data Science Platforms: A strategic initiative to build advanced artificial intelligence and data science capabilities across drug discovery, clinical development, and operational analytics, integrating large datasets with machine learning to accelerate medicine development and improve decision‑making.
    •	Nerve Live Data Insight Platform: A digital transformation program combining a data lake, analytical engine, and application modules to unify operational data and generate real‑time insights for managing and optimizing clinical trials, R&D processes, and business operations.
    •	Novartis Biome Innovation Hubs: A network of global innovation hubs that partners with external technology and healthcare organizations to co‑create and scale digital health solutions, including mobile health tools and point‑of‑care diagnostic innovations aimed at expanding patient reach and improving healthcare delivery.
	 * Never say : played the key role, or lead ... or mentor other guys.
   2)Experience making rule
 - I gave you project name or scope or scenario for every company above.
  In Next step, while relying on the project scope and company overview which I gave you above , also you must try to  tailore to JD so that I could be best candidate for this position(passing over 95% ATS).  
   plz read JD carefully and understand what responsibility(and ability)  this role require me to have. then try to make resume so that I have perfectly suitable ability for this role. 
   I mean plz try to combine my project scope or scenario with JD's role sealessly very very naturally.
   (JD's requirement <- Balance -> Project Scope)
 - If role is backend, or frontend or fullstack or data engineering, make relevant experience solely for corresponding field
   (I mean JD's role is backend, then must not mention frontend, also if JD mention data engineernig, don't mention others, but only mention data engineering
    Please comply with this rule for all companies) 
 - Sometimes, job description does not include enough skills and experience and they mention only a few core  skills or experience rather than listing many.
   So plz think what additional skills and experience would be good to be added and plz add them in experience section. 
   (ex: React is mentioned in JD, but Next.js is not mentioned. in this case You should mention Next.js in skill section too. This is vivid example of mentioning related skills.)
 - make 7 bullet for Fresenius Medical Care, 8 bullet work bio for Capgemini, 7 Novartis.
 - as long as possible technically, try to mention many core and important skills and keywords mentioned in JD when making experience section bullet points.
 - Never make technically fake sentences.
   Never say more than one languages for one company project(ex: Used Java and C# for....: I hate this and only use one language for one company project).
 - When describing the experience, must focus about 70% on the primary programming languages, technology stacks, libraries, frameworks, and tools listed in the job description, and use the remaining 30% to highlight other related relevant technologies or skills that strengthen my overall profile. 
   You should mention core language or framework related another libraries or framework too(ex: JD mentions only React, but you should mention also React Native, Next.js. This does not mean you mention React, React Native, Next.js for all companies. plz align them with distributed companies. 
   This is just one example, another->if JD mention Flask, then ok to mention FastAPI, so many examples for this, right? If Java mentioned->Kotlin ok,......)
 - hightlight(bold font) as many skills as possible in work experience to showcase my experience of I have already used them.
 - end each bullet with fullstop.
 - Never use "Led ..." if you want to emphasize that kind of leadership skill, use "Played the key role in ... "
   Never use these expressions: "scalable", "high-performance".                
 3)Quantifiable Impact - Non-Negotiable:
  * MUST add quantifiable results in 2–3 technical bullet points per company, using strict numeric expressions (e.g., “around X records,” “over X users,” “approximately X requests per day”).
    * **NEVER show quantifiable impacts in every bullet point**. Only 1 -2 bullets should feature quantifiable metrics that reflect the most significant impact of your work. The other bullets should focus on describing actions, responsibilities, or broader context without quantification.
  * Quantitative values should be determined based on the company stage and available data. For early-stage or smaller companies, provide realistic approximations of scale. For more established companies, use exact numbers or well-supported approximations based on available metrics.
  * In most cases, if precise figures are unavailable, MUST use approximations like “around,” “over,” or “approximately” to convey a sense of scale or impact.
    * Do not use vague, non-numeric descriptors or qualitative terms.
  * Quantification may include:
    * Exact numeric volumes (e.g., records, requests)
    * Numeric user counts (e.g., user base, active users)
    * Numeric timeframes (e.g., release cycles, project duration)
    * Team or system scales (e.g., number of services, team size)
    * Workload or throughput values
    * Deployment frequencies (e.g., updates, releases per time period)
  * MUST NOT use:
    * Textual numeric expressions (e.g., *thousands, hundreds, tens of thousands*)
    * Plus symbols (e.g., *X+ users*)
    * Percentages in any form
    * Vague, non-numeric descriptors (e.g., *a lot of*, *many*, *significant*).
  **Important (Required for All Bullet Points)**: NEVER use percentage expressions like "x%", MUST remove or replace it with numeric expressions.

Please make html in new code panel with following styles:
I give you sample html code for one resume below:
                   
<!DOCTYPE html>

<html lang="en">



<head>

  <meta charset="UTF-8">

  <meta name="viewport" content="width=device-width, initial-scale=1.0">

  <style>

    body {

      font-family: Arial, Helvetica, sans-serif;

    }
    
    p {
      padding: 0;
      margin: 0;        
      line-height: 1.4;
    }
     

    .title {

      font-size: 16px;

      color: black;

    }



    .role {

      font-size: 14px;

    }



    .section-title {

      font-size: 14px;

      font-weight: bold;

      color: #ff6c00;

      display: inline-block;

      padding-bottom: 5px;

    }



    .underline {

      width: 100%;

      height: 1px;

      background-color: black;

    }



    .content {

      font-size: 12px;

    }



    .bold {

      font-weight: bold;
      margin: 5px 0 5px 0;

    }



    ul {

      list-style-type: disc;

      padding-left: 20px;
      margin: 0;

    }
                   
    ul li {
      margin-bottom: 0px;     
      line-height: 1.4;      
    }                   

    .education-content {

      font-weight: bold;
      display: block;
      margin: 5px 0 5px 0;

    }

  </style>

</head>



<body>



  <!-- Header Section -->

  <div class="title">Patryk Pham</div>

  <div class="role">Senior Software Engineer</div>

  <div class="content" style="margin-top: 4px;">patryk.p0412@gmail.com • Cracow, Poland • www.linkedin.com/in/patryk-p-2080a73b3/</div>



  <!-- Summary Section -->

  <div class="section-title" style="margin-top: 5px;">Summary</div>

  <div class="underline"></div>

  <div class="content">

    <p>Experienced JD's job title ....</p>

  </div>



  <!-- Skills Section -->

  <div class="section-title">Skills</div>

  <div class="underline"></div>

  <div class="content">

    <ul>

      <li><strong>[Category1]:</strong>[skill1], [skill2], [skill3], ...</li>

      <li><strong>[Category2]:</strong>[skill1], [skill2], [skill3], ...</li>
      (....)
    </ul>

  </div>



  <!-- Work Experience Section -->

  <div class="section-title">Work Experience</div>

  <div class="underline"></div>

  <div class="content">

    <!-- Fresenius Medical Care -->

    <div class="bold">Fresenius Medical Care • Senior Data Engineer • 06/2023 - Present</div>

    <ul>
      <li>[experience1.... <strong>${core skill},.... ${core skill}...]</strong></li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience3.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...
     </ul>



    <!-- Capgemini -->

    <div class="bold">Capgemini • Data Engineer • 09/2020 - 05/2023</div>

    <ul>

      <li>[experience1.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...
    </ul>



    <!-- Novartis -->

    <div class="bold">Novartis • Junior Software Developer • 05/2018 - 08/2020</div>

    <ul>

      <li>[experience1.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience3.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...

    </ul>

  </div>



  <!-- Education Section -->

  <div class="section-title">Education</div>

  <div class="underline"></div>

  <div class="content">

    <ul>

      <li><span class="education-content">Babes-Bolyai University • Bachelor's Degree in COMPUTER AND INFORMATION SCIENCES AND SUPPORT SERVICES
 • 2013 - 2017</span></li>

    </ul>

  </div>



</body>



</html>
(Don't add another margin or padding in your decision)   

All in all, you must give me in code panle for resume so that I can copy.         
But before give me resume, gpt, plz judge whether job description say whether this job require me to work hybrid or on site.
If this is not fully remote, stop making resume and let me know the reason.(I am based in Romania).
If full remote, then go to next step.

Next step: As I emphasize much about Capgemini, you must make very very specific experience section not generic sound one. but you always forget this critical rule cuz you are foolish.
So before making html resume, let me know your plan of project name & project scenario for Capgemini briefly(Think you are presenting your project idea in very detail with real word usage example). 
Here, important thing is that one engineer can not do everything for one project development, so never say that I participated in so many works. all bullets of experience must be specific while related to project name so you must make  hiring manger 
think, oh it is right and reasonable that this guy has worked on this project, not fake engineer. make sense? make sense? 

I said you must use strong tag in html to make bold text not --, so after you add plan, plz confirm you will only use strong tag for bold text.
If I agree with you, 
then, start to make resume.                   
""")
    pyautogui.hotkey('ctrl', 'v')
    pyautogui.press('enter')
    print("Text sent to ChatGPT. Save DOC output manually as 'latest.docx', then press F3.")


def copy_selected_html_to_resume():
    """F3: Copy selected HTML code (resume), create resume.html, generate resume.pdf"""
    # Copy selected HTML from clipboard
    keyboard.press_and_release('ctrl+c')
    time.sleep(0.2)
    html_text = pyperclip.paste()
    if not html_text.strip():
        print("No HTML selected. Please select your HTML code first.")
        return
    
    # Append to save.docx
    save_to_word(html_text)
    print(f"Appended selected HTML to {SAVE_FILE}")
    
    # Create new resume.html (overwrite old if exists)
    if os.path.exists(HTML_FILE):
        os.remove(HTML_FILE)
    with open(HTML_FILE, "w", encoding="utf-8") as f:
        f.write(html_text)
    print(f"Created new {HTML_FILE}")

    # Generate PDF from resume.html
    if os.path.exists(PDF_FILE):
        os.remove(PDF_FILE)
	


    HTML(HTML_FILE, encoding="utf-8").write_pdf(
        PDF_FILE,
        stylesheets=[
            CSS(string="""
                @page {
                    size: A4;
                    margin-top: 10mm;
                    margin-bottom: 7mm;
                    margin-left: 5mm;
                    margin-right: 5mm;
                }
                body {
                    font-family: "Arial", sans-serif;
                    letter-spacing: 0.02em;
                }
            """)
        ]
    )

    os.startfile(PDF_FILE)
    print("Generated resume PDF!")

# ================= HOTKEYS =================
keyboard.add_hotkey('[', send_to_chatgpt)  # Press `[` to send Job Description to ChatGPT
keyboard.add_hotkey(']', copy_selected_html_to_resume)  # Press `]` to generate Resume PDF

print("[: Send selected job description to ChatGPT")
print("]: Generate Resume PDF from HTML")
keyboard.wait()
