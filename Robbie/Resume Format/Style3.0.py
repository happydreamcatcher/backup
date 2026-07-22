import os
from datetime import datetime
import time
import pyperclip
import pyautogui
import keyboard
from weasyprint import HTML, CSS
from docx import Document

# ================= CONFIG =================
SAVE_FILE = "save.docx"
PDF_FILE = "MichalTruongCV.pdf"
HTML_FILE = "CV.html"
GPT_APP_TITLE = "ChatGPT"  # Adjust to your app window title
PROMPT_TEXT = "Please improve this text: "

# ==========================

def save_to_word(text):
    """Append text to a Word file"""
    today = datetime.now().strftime("%m-%d")
    filename = "Description\\" + today + ".docx"
    if os.path.exists(filename):
        doc = Document(filename)
    else:
        doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)
    print(f"Saved to {filename}")


def send_to_chatgpt():
    """F2: Copy selected text, send it to ChatGPT for generating resume"""
    # Copy selected text (Job Description)
    keyboard.press_and_release('ctrl+c')
    time.sleep(0.2)
    text = pyperclip.paste()
    if not text.strip():
        print("No text selected. Please select text first.")
        return

    save_to_word(text)

    pyautogui.keyDown('ctrl')
    pyautogui.press('tab')
    pyautogui.keyUp('ctrl')

    # Open new chat
    pyautogui.hotkey('ctrl', 'shift', 'o')
    time.sleep(0.5)

    # Paste prompt + text for ChatGPT
    pyperclip.copy(text + """

Above is job description and following is current resume and I am tailoring my resume according to job description.
And I also want to make resume.
1. Resume
I am gonnna make resume in order to bid on jobs.
Here is my resume skeleton and rule of how to structure and make resume.
------* Resume Skeleton --------
0. Head
Michal Truong
Senior Software Engineer
michal.tr88@hotmail.com • Gdańsk, Poland • www.linkedin.com/in/michal-truong-58b7093b6/

1. Summary
2. Skills
3. Experience
- HoNa Health • Senior Software Engineer • 04/2023 - Present
- Endava	• Senior Software Engineer • 03/2020 - 02/2023		              
- Miquido • Junior Software Developer • 04/2017 - 12/2019 	              
- Self Employed • Independent Coding Coach • 07/2016 - 04/2017                      
4. Education
  Ho Chi Minh City University of Technology • Bachelor's Degree in Computational and Applied Mathematics • Vietnam • 2012 - 2016 		  

------* Resume make rules --------
1. Summary Section
Write a strong professional summary(10 years).
Start with "Experienced Senior Software Engineer..."
In 4 concise sentences, summarize my experience, core strengths, and what makes me a good fit for the role. 
Showcase to my extensive experience including the project type and project names briefly that I have been participated in last three companies(HoNa Health, Endava, Miquido: Never mention company names).
When next sentence starts in same line with previous sentence, leave 4 letters gap between sentences in summary.
 
2. Core Skills Section
 - At least, skill section must include almost tech stack or skills of JD.
   Excepted for JD mentioned Skills, also, plz  understand what related skills are essential or necessary and, mention them even though they are not stated in JD.
 - (skill section should include 25 - 35 skills)
 - under category, give me following visual way:
   Category1 Name(bold font): Skill1, Skill2(all in one line)
   Category1 Name(bold font): Skill1, Skill2((all in one line))
   Use bullet point for every category.
   don't hightlight keywords at all.
 - The maximum category numbers are 8. Never make more than 8 categories.
3. Experience Section
  1)Project overview & Scope in every company 
   - HoNa Health: Senior Role
   * HoNa Health is a healthcare technology startup focused on helping providers make sense of fragmented patient data. Its platform integrates with electronic health record systems, retrieves broader patient history, and turns large volumes of medical records into concise, customizable clinical summaries that fit directly into provider workflows.
   * Work Scope:  
     •	Clinical Summary & Visit-Prep Platform: Worked as part of the development team on HoNa Health’s provider-facing platform that gathered patient records from EHR systems and external data sources, then organized them into specialty-specific, easy-to-review summaries for care teams before appointments. Contributed to building responsive interfaces and workflow-driven features that helped clinicians review labs, notes, medications, and source-linked history more efficiently.
     •	Healthcare Data Integration & Customization Features: Contributed to product features that supported medical-record retrieval, multi-source data integration, and customizable summary templates, allowing providers to filter and view the most relevant patient information based on their specialty and visit context. Worked with the team on web features that improved navigation of longitudinal patient data and reduced the number of clicks needed to access key insights.
   - In Endava: Senior Role
	 * Endava is software dev company. so I can work various type of projects. After reading JD, make very specific project name(less popular and small project) in a related industry and definitely different project with very reasonable sceario.
     Here, most important thing is that project name must not sound like generic. It must sound real world project name and scenario too. I hate ("IAM system for Sass platform") like generic project name. When hiring manager sees the project name, he must understand oh what this project is for.
     Also, project name must not include platform, it sounds awkward, make sense?
   - Miquido: Junior Role
   * Miquido is a software development company that designs and builds custom digital products for businesses across industries such as fintech, healthcare, e-commerce, travel, and entertainment. The company is known for delivering web applications, mobile apps, product design, and AI-powered solutions for clients ranging from startups to enterprise brands.
   * Work Scope:  
     •	Policyholder Self-Service & Claims Intake Platform: Contributed to the development of a web platform for an insurance client that enabled policy viewing, digital claims submission, document upload, claim-status tracking, and billing/payment-related interactions, helping streamline claims processing and improve customer self-service.
     •	Omnichannel Quote Distribution & Product Personalization Module: Worked on customer-facing web features for an insurance client that supported omnichannel policy distribution, workflow automation, and personalized insurance offer experiences based on user and product data.
	 * Never say : played the key role, or lead ... or mentor other guys. I started as software intern then promoted to software developer in Miquido.
  - Self Employed
   * I have worked as coding coach to teach several students programming and basic software engineering principles. So use following 3 bullets directly in experience section.
   * Work Scope:  
     •	Taught students core programming concepts, including problem-solving, data structures, algorithms, and debugging techniques, through personalized one-on-one coaching.
     •	Introduced basic software engineering principles such as code organization, version control, testing, and writing clean, maintainable code.
     •	Designed tailored lesson plans and practice exercises to match individual learning goals, helping students improve technical confidence and coding proficiency.
  2)Experience making rule
 - I gave you project name or scope or scenario for every company above.
  In Next step, while relying on the project scope and company overview which I gave you above , also you must try to  tailore to JD so that I could be best candidate for this position(passing over 95% ATS).  
   plz read JD carefully and understand what responsibility(and ability)  this role require me to have. then try to make resume so that I have perfectly suitable ability for this role. 
   I mean plz try to combine my project scope or scenario with JD's role sealessly very very naturally.
   (JD's requirement <- Balance -> Project Scope)
 - If role is backend, or frontend or fullstack or data engineering, make relevant experience solely for corresponding field
   (I mean JD's role is backend, then must not mention frontend, also if JD mention data engineernig, don't mention others, but only mention data engineering
    Please comply with this rule for all companies) 
 - Sometimes, job description does not include enough skills and experience and they mention only a few core  skills or experience rather than listing many.
   So plz think what additional skills and experience would be good to be added and plz add them in experience section. 
 - make 8 for HoNa Health, 8 bullet work bio for Endava, 7 Miquido.
 - as long as possible technically, try to mention many core and important skills and keywords mentioned in JD when making experience section bullet points.
 - Never make technically fake sentences.
 - When describing the experience, must focus about 70% on the primary programming languages, technology stacks, libraries, frameworks, and tools listed in the job description, and use the remaining 30% to highlight other related relevant technologies or skills that strengthen my overall profile. 
   You should mention core language or framework related another libraries or framework too.
 - hightlight(bold font) as many skills as possible in work experience to showcase my experience of I have already used them.
 - end each bullet with fullstop.
 - Never use "Led ..." if you want to emphasize that kind of leadership skill, use "Played the key role in ... "
   Never use these expressions: "scalable", "high-performance".                
 3)Quantifiable Impact - Non-Negotiable:
  * MUST add quantifiable results in 2–3 technical bullet points per company, using strict numeric expressions (e.g., “around X records,” “over X users,” “approximately X requests per day”).
    * **NEVER show quantifiable impacts in every bullet point**. Only 1 -2 bullets should feature quantifiable metrics that reflect the most significant impact of your work. The other bullets should focus on describing actions, responsibilities, or broader context without quantification.
  * Quantitative values should be determined based on the company stage and available data. For early-stage or smaller companies, provide realistic approximations of scale. For more established companies, use exact numbers or well-supported approximations based on available metrics.
  * In most cases, if precise figures are unavailable, MUST use approximations like “around,” “over,” or “approximately” to convey a sense of scale or impact.
    * Do not use vague, non-numeric descriptors or qualitative terms.
  * Quantification may include:
    * Exact numeric volumes (e.g., records, requests)
    * Numeric user counts (e.g., user base, active users)
    * Numeric timeframes (e.g., release cycles, project duration)
    * Team or system scales (e.g., number of services, team size)
    * Workload or throughput values
    * Deployment frequencies (e.g., updates, releases per time period)
  * MUST NOT use:
    * Textual numeric expressions (e.g., *thousands, hundreds, tens of thousands*)
    * Plus symbols (e.g., *X+ users*)
    * Percentages in any form
    * Vague, non-numeric descriptors (e.g., *a lot of*, *many*, *significant*).
  **Important (Required for All Bullet Points)**: NEVER use percentage expressions like "x%", MUST remove or replace it with numeric expressions.

Please make html in new code panel with following styles:
I give you sample html code for one resume below:
                   
<!DOCTYPE html>

<html lang="en">



<head>

  <meta charset="UTF-8">

  <meta name="viewport" content="width=device-width, initial-scale=1.0">

  <style>

    body {

      font-family: Georgia;

    }
    
    p {
      padding: 0;
      margin: 0;        
      line-height: 1.4;
    }
     

    .title {

      font-size: 16px;

      color: black;

    }



    .role {

      font-size: 14px;

    }



    .section-title {

      font-size: 14px;

      font-weight: bold;

      color: #ff6c00;

      display: inline-block;

      padding-bottom: 5px;

    }



    .underline {

      width: 100%;

      height: 2px;

      background-color: black;

    }



    .content {

      font-size: 12px;

    }



    .bold {

      font-weight: bold;
      margin: 5px 0 5px 0;

    }



    ul {

      list-style-type: disc;

      padding-left: 20px;
      margin: 0;

    }
                   
    ul li {
      margin-bottom: 0px;     
      line-height: 1.4;      
    }                   

    .education-content {

      font-weight: bold;
      display: block;
      margin: 5px 0 5px 0;

    }

  </style>

</head>



<body>



  <!-- Header Section -->

  <div class="title">Michal Truong</div>

  <div class="role">Senior Software Engineer</div>

  <ul class="content" style="margin-top: 4px;"><li>michal.tr88@hotmail.com</li><li>Gdańsk, Poland</li><li>www.linkedin.com/in/michal-truong-58b7093b6/</li></ul>



  <!-- Summary Section -->

  <div class="section-title" style="margin-top: 5px;">Summary</div>

  <div class="underline"></div>

  <div class="content">

    <p>Experienced JD's job title ....</p>

  </div>



  <!-- Skills Section -->

  <div class="section-title">Skills</div>

  <div class="underline"></div>

  <div class="content">

    <ul>

      <li><strong>[Category1]:</strong>[skill1], [skill2], [skill3], ...</li>

      <li><strong>[Category2]:</strong>[skill1], [skill2], [skill3], ...</li>
      (....)
    </ul>

  </div>



  <!-- Work Experience Section -->

  <div class="section-title">Work Experience</div>

  <div class="underline"></div>

  <div class="content">

    <!-- HoNa Health -->

    <div class="bold">HoNa Health • Senior Software Engineer • Poland • 04/2023 - Present</div>

    <ul>
      <li>[experience1.... <strong>${core skill},.... ${core skill}...]</strong></li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience3.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...
     </ul>



    <!-- Endava -->

    <div class="bold">Endava • Senior Software Engineer • Poland • 03/2020 - 02/2023</div>

    <ul>

      <li>[experience1.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...
    </ul>



    <!-- Miquido -->

    <div class="bold">Miquido • Junior Software Developer • Poland • 04/2017 - 12/2019</div>

    <ul>

      <li>[experience1.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience3.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...

    </ul>
    <div class="bold">Self Employed • Independent Coding Coach • Vietnam • 07/2016 - 04/2017</div>
    <ul>

      <li>[experience1.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience2.... <strong>${core skill},.... ${core skill}</strong>...]</li>

      <li>[experience3.... <strong>${core skill},.... ${core skill}</strong>...]</li>
      ...

    </ul>

  </div>

  <!-- Education Section -->

  <div class="section-title">Education</div>

  <div class="underline"></div>

  <div class="content">

    <ul>

      <li><span class="education-content">Ho Chi Minh City University of Technology • Bachelor's Degree • Computational and Applied Mathematics • Ho Chi Minh City, Vietnam • 2012 - 2016</span></li>

    </ul>

  </div>



</body>



</html>
(Don't add another margin or padding in your decision)   

All in all, you must give me in code panle for resume so that I can copy.         
But before give me resume, gpt, plz judge whether job description say whether this job require me to work hybrid or on site.
If this is not fully remote, stop making resume and let me know the reason.(I am based in Poland).
If full remote, then go to next step.

Next step: As I emphasize much about Endava, you must make very very specific experience section not generic sound one. but you always forget this critical rule cuz you are foolish.
So before making html resume, let me know your plan of project name & project scenario for  briefly(Think you are presenting your project idea in very detail with real word usage example). 
Here, important thing is that one engineer can not do everything for one project development, so never say that I participated in so many works. all bullets of experience must be specific while related to project name so you must make  hiring manger 
think, oh it is right and reasonable that this guy has worked on this project, not fake engineer. make sense? make sense? 

After briefly showing your plan for Endava, also briefly mention what languages & frameworks you are gonna use for another companies briefly, Avoid to mention skills that are not explicitly mentioned in JD.
Also, I said you must use strong tag in html to make bold text not --, so after you add plan, plz confirm you will only use strong tag for bold text.
If I agree with you, 
then, start to make resume.                   
""")
    pyautogui.hotkey('ctrl', 'v')
    pyautogui.press('enter')
    print("Text sent to ChatGPT. Save DOC output manually as 'latest.docx', then press F3.")


def copy_selected_html_to_resume():
    """F3: Copy selected HTML code (resume), create resume.html, generate resume.pdf"""
    # Copy selected HTML from clipboard
    keyboard.press_and_release('ctrl+c')
    time.sleep(0.2)
    html_text = pyperclip.paste()
    if not html_text.strip():
        print("No HTML selected. Please select your HTML code first.")
        return
    
    # Append to save.docx
    save_to_word(html_text)
    print(f"Appended selected HTML to {SAVE_FILE}")
    
    # Create new resume.html (overwrite old if exists)
    if os.path.exists(HTML_FILE):
        os.remove(HTML_FILE)
    with open(HTML_FILE, "w", encoding="utf-8") as f:
        f.write(html_text)
    print(f"Created new {HTML_FILE}")

    # Generate PDF from resume.html
    if os.path.exists(PDF_FILE):
        os.remove(PDF_FILE)
	


    HTML(HTML_FILE, encoding="utf-8").write_pdf(
        PDF_FILE,
        stylesheets=[
            CSS(string="""
                @page {
                    size: A4;
                    margin-top: 10mm;
                    margin-bottom: 7mm;
                    margin-left: 5mm;
                    margin-right: 5mm;
                }
                body {
                    font-family: "Georgia", sans-serif;
                    letter-spacing: 0.02em;
                }
            """)
        ]
    )

    os.startfile(PDF_FILE)
    print("Generated resume PDF!")

# ================= HOTKEYS =================
keyboard.add_hotkey('[', send_to_chatgpt)  # Press `[` to send Job Description to ChatGPT
keyboard.add_hotkey(']', copy_selected_html_to_resume)  # Press `]` to generate Resume PDF

print("[: Send selected job description to ChatGPT")
print("]: Generate Resume PDF from HTML")
keyboard.wait()
